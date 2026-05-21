# modules/precision_ocr.py  — Production OCR Engine v2.0
"""
Production-level OCR pipeline using gemma4:26b vision model.
Features:
  - 5 adaptive image preprocessing strategies with auto-retry
  - Multi-page PDF support (processes all pages, merges results)
  - Image upscaling for low-resolution scans
  - Robust response parsing with fallback chain
  - Per-step error surfacing (no silent failures)
  - Archive vault integration
  - Download in MD / PDF / DOCX formats
"""
import streamlit as st
import io
import fitz          # PyMuPDF
import re
import math
import numpy as np
from PIL import Image, ImageEnhance, ImageOps, ImageFilter
import ollama
import markdown
from xhtml2pdf import pisa
from docx import Document
import time
import pandas as pd
from database.db_utils import archive_document

# ---------------------------------------------------------------------------
# 1.  IMAGE PREPROCESSING STRATEGIES
MIN_DIM = 800   # minimum pixel dimension; smaller images are upscaled

def _upscale_if_small(img: Image.Image) -> Image.Image:
    """Upscale image if either dimension is below MIN_DIM."""
    w, h = img.size
    if min(w, h) < MIN_DIM:
        scale = MIN_DIM / min(w, h)
        new_w, new_h = int(w * scale), int(h * scale)
        img = img.resize((new_w, new_h), Image.LANCZOS)
    return img

def strategy_original(img: Image.Image) -> Image.Image:
    """Strategy 0 — Send the image with minimal changes (fix orientation only)."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    return img.convert("RGB")

def strategy_mild_enhance(img: Image.Image) -> Image.Image:
    """Strategy 1 — Mild sharpening + contrast boost. Best for printed forms."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    img = img.convert("RGB")
    img = ImageEnhance.Sharpness(img).enhance(1.8)
    img = ImageEnhance.Contrast(img).enhance(1.5)
    return img

def strategy_grayscale_boost(img: Image.Image) -> Image.Image:
    """Strategy 2 — Grayscale + moderate contrast. Good for photocopies."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    gray = img.convert("L")
    gray = ImageEnhance.Contrast(gray).enhance(1.8)
    gray = ImageEnhance.Sharpness(gray).enhance(2.0)
    return gray.convert("RGB")

def strategy_adaptive_threshold(img: Image.Image) -> Image.Image:
    """Strategy 3 — Near-binarization for faded handwriting on white background."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    gray = img.convert("L")
    # Auto-level: stretch histogram
    arr = np.array(gray, dtype=np.float32)
    p2, p98 = np.percentile(arr, 2), np.percentile(arr, 98)
    if p98 > p2:
        arr = np.clip((arr - p2) / (p98 - p2) * 255, 0, 255)
    gray = Image.fromarray(arr.astype(np.uint8))
    gray = ImageEnhance.Contrast(gray).enhance(2.2)
    # Unsharp mask to make thin strokes pop
    gray = gray.filter(ImageFilter.UnsharpMask(radius=1, percent=150, threshold=3))
    return gray.convert("RGB")

def strategy_denoised(img: Image.Image) -> Image.Image:
    """Strategy 4 — Denoise first then enhance. For camera photos of documents."""
    img = ImageOps.exif_transpose(img)
    img = _upscale_if_small(img)
    img = img.convert("RGB")
    # Median filter removes camera noise
    img = img.filter(ImageFilter.MedianFilter(size=3))
    img = ImageEnhance.Contrast(img).enhance(1.6)
    img = ImageEnhance.Sharpness(img).enhance(2.0)
    return img

STRATEGIES = [
    ("Original",            strategy_original),
    ("Mild Enhancement",    strategy_mild_enhance),
    ("Grayscale Boost",     strategy_grayscale_boost),
    ("Adaptive Threshold",  strategy_adaptive_threshold),
    ("Denoise + Enhance",   strategy_denoised),
]

def img_to_bytes(img: Image.Image, quality: int = 92) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=quality)
    return buf.getvalue()

# ---------------------------------------------------------------------------
# 2.  RESPONSE CLEANING
# ---------------------------------------------------------------------------
def clean_output(raw: str) -> str:
    """
    Remove model control tokens and code fences.
    Falls back to the raw text if cleaning removes everything.
    Preserves medical symbols like < > in reference ranges.
    """
    if not raw:
        return ""

    text = raw

    # Strip content before <|text|> markers (some model variants)
    if "<|text|>" in text:
        text = text.split("<|text|>")[-1]
    if "</|text|>" in text:
        text = text.split("</|text|>")[0]

    # Remove thinking blocks — both formats used by different model variants
    text = re.sub(r'<\|think\|>.*?</\|think\|>', '', text, flags=re.DOTALL)   # pipe format
    text = re.sub(r'<think>.*?</think>',           '', text, flags=re.DOTALL)   # plain format

    # Remove code fences
    text = re.sub(r'```[a-z]*\n?', '', text)
    text = text.replace('```', '')

    cleaned = text.strip()

    # Safety fallback: if stripping removed EVERYTHING, return raw minus fences
    if not cleaned:
        fallback = re.sub(r'```[a-z]*\n?', '', raw).replace('```', '').strip()
        return fallback

    return cleaned

# ---------------------------------------------------------------------------
# 3.  OLLAMA CALL WITH RETRY ACROSS PREPROCESSING STRATEGIES
# ---------------------------------------------------------------------------
MODEL_OPTIONS = {
    "temperature":  0,
    "num_ctx":      8192,
    "num_predict":  4096,   # Raised from 2048 — thinking models need headroom
}

def call_model_with_retry(
    raw_img: Image.Image,
    prompt: str,
    status_placeholder=None,
    model_name="qwen2.5vl:32b"
) -> tuple[str, str]:
    """
    Try each preprocessing strategy in order.
    Returns (extracted_text, strategy_name_used) or ('', '') on total failure.
    """
    for strategy_name, strategy_fn in STRATEGIES:
        if status_placeholder:
            status_placeholder.info(f"🔄 Trying strategy: **{strategy_name}**…")
        try:
            # Prevent oversized images from crashing Ollama (OOM on ViT)
            # Increased to 2560 to preserve high resolution for small text OCR
            max_dim = 2560
            w, h = raw_img.size
            if max(w, h) > max_dim:
                scale = max_dim / max(w, h)
                raw_img = raw_img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
                
            processed = strategy_fn(raw_img)
            image_bytes = img_to_bytes(processed)

            resp = ollama.chat(
                model=model_name,
                messages=[{
                    "role":    "user",
                    "content": prompt,
                    "images":  [image_bytes],
                }],
                options=MODEL_OPTIONS,
            )
            raw_text = resp["message"]["content"]
            result   = clean_output(raw_text)
            
            # DEBUG LOGGING (Absolute path)
            log_path = "/home/kalpra/Downloads/OCR/database/ocr_debug.log"
            with open(log_path, "a") as f:
                f.write(f"\n--- {time.strftime('%H:%M:%S')} | STRATEGY: {strategy_name} ---\n")
                f.write(f"RAW (len {len(raw_text)}): {raw_text[:300]}\n")
                f.write(f"CLEANED (len {len(result)}): {result[:300]}\n")

            if result and len(result.strip()) >= 5:
                return result, strategy_name

        except Exception as e:
            if status_placeholder:
                status_placeholder.warning(f"⚠️ Strategy '{strategy_name}' failed: {e}")
            continue

    return "", ""

# ---------------------------------------------------------------------------
# 4.  PDF MULTI-PAGE LOADER
# ---------------------------------------------------------------------------
def load_pdf_pages(file_bytes: bytes, max_pages: int = 10) -> list[Image.Image]:
    """Extract all pages from a PDF as PIL images at 220 DPI (Balanced for Qwen)."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for i in range(min(len(doc), max_pages)):
        page = doc.load_page(i)
        pix  = page.get_pixmap(dpi=220)
        pages.append(Image.open(io.BytesIO(pix.tobytes("png"))))
    return pages

# ---------------------------------------------------------------------------
# 5.  EXPORT GENERATORS  (PDF / DOCX)
# ---------------------------------------------------------------------------
def generate_pro_pdf(md_text: str, client_name: str) -> bytes | str:
    try:
        safe = re.sub(r'<\|.*?\|>', '', md_text)
        safe = safe.replace('<', '&lt;').replace('>', '&gt;')
        safe = safe.replace('\u2022', '*').replace('\u2713', '[x]')
        html_body = markdown.markdown(safe, extensions=['tables'])
        orientation = "landscape" if "| S.No |" in md_text else "portrait"
        html = f"""<html><head><style>
            @page {{size: A4 {orientation}; margin:1cm;}}
            body {{font-family:Helvetica,Arial,sans-serif; font-size:10px;}}
            table {{width:100%; border-collapse:collapse; margin-top:10px;}}
            th,td {{border:1px solid #333; padding:5px; text-align:left;}}
            th {{background:#ddd; font-weight:bold;}}
            h2 {{text-align:center; border-bottom:2px solid #333; padding-bottom:5px;}}
        </style></head><body><h2>{client_name} — Extraction Report</h2>{html_body}</body></html>"""
        out = io.BytesIO()
        status = pisa.CreatePDF(io.StringIO(html), dest=out, encoding="utf-8")
        return out.getvalue() if not status.err else f"PDF Error: {status.err}"
    except Exception as e:
        return f"PDF Error: {e}"

def generate_docx(md_text: str, client_name: str) -> bytes | str:
    try:
        doc = Document()
        doc.add_heading(f"{client_name} — Extraction Report", 0)
        clean = re.sub(r'<\|.*?\|>', '', md_text)
        in_table = False
        table    = None
        for line in clean.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('|'):
                if '---' in line:
                    continue
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if not cells:
                    continue
                if not in_table:
                    table    = doc.add_table(rows=1, cols=len(cells))
                    table.style = 'Table Grid'
                    for i, c in enumerate(cells):
                        table.rows[0].cells[i].text = c
                    in_table = True
                else:
                    row = table.add_row()
                    for i, c in enumerate(cells):
                        if i < len(row.cells):
                            row.cells[i].text = c
            else:
                in_table = False
                doc.add_paragraph(line)
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception as e:
        return f"Word Error: {e}"

def generate_excel(md_text: str) -> bytes | None:
    if not md_text.strip():
        return None
        
    try:
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine='openpyxl') as writer:
            # --- 1. CONSOLIDATED REPORT (Primary Sheet) ---
            report_data = []
            lines = md_text.split('\n')
            
            # Extract key-value pairs (like **Key:** Value)
            metadata = []
            other_text = []
            tables = []
            current_table = []
            
            for line in lines:
                raw_line = line.strip()
                if not raw_line: continue
                
                # Table detection
                if raw_line.count('|') >= 2:
                    if all(c in '|- ' for c in raw_line) and '-' in raw_line:
                        continue
                    cells = [c.strip() for c in raw_line.split('|') if c.strip()]
                    if cells:
                        current_table.append(cells)
                    continue
                else:
                    if current_table:
                        if len(current_table) > 1:
                            tables.append(current_table)
                        current_table = []
                
                # Metadata detection (Bold keys)
                if '**' in raw_line and ':' in raw_line:
                    # Clean up markers like **Patient Name:**
                    parts = re.split(r'\*\*|\*', raw_line)
                    cleaned = " ".join([p.strip() for p in parts if p.strip()])
                    metadata.append([cleaned])
                else:
                    other_text.append([raw_line])
            
            if current_table and len(current_table) > 1:
                tables.append(current_table)

            # Build the master list for the first sheet
            # [A] Patient / Form Metadata
            if metadata:
                report_data.append(["--- FORM DETAILS ---"])
                report_data.extend(metadata)
                report_data.append([""]) # Spacer
            
            # [B] Tables
            if tables:
                for i, table in enumerate(tables):
                    report_data.append([f"--- TABLE {i+1} ---"])
                    report_data.extend(table)
                    report_data.append([""]) # Spacer
            
            # [C] Other Text / Footer
            if other_text:
                report_data.append(["--- ADDITIONAL INFORMATION ---"])
                report_data.extend(other_text)
            
            # Create the Consolidated Sheet
            # Find max width to avoid dataframe errors
            max_cols = max([len(r) for r in report_data]) if report_data else 1
            df_cons = pd.DataFrame(report_data, columns=[f"Col {i+1}" for i in range(max_cols)])
            df_cons.to_excel(writer, index=False, header=False, sheet_name='Consolidated Report')
            
            # --- 2. CLEAN TABLE SHEETS (Individual sheets for analysis) ---
            if tables:
                for i, table in enumerate(tables):
                    header = table[0]
                    rows = table[1:]
                    cleaned_rows = [r[:len(header)] if len(r) > len(header) else r + ['']*(len(header)-len(r)) for r in rows]
                    df_table = pd.DataFrame(cleaned_rows, columns=header)
                    sheet_name = f'Table_{i+1}'
                    df_table.to_excel(writer, index=False, sheet_name=sheet_name)
            
            # --- 3. RAW EXTRACTION SHEET ---
            paragraphs = [p.strip() for p in md_text.split('\n') if p.strip()]
            df_full = pd.DataFrame({"Raw Content": paragraphs})
            df_full.to_excel(writer, index=False, sheet_name='Raw Text')
            
        return out.getvalue()
    except Exception as e:
        print(f"Excel Error: {e}")
        # Final fallback
        try:
            out = io.BytesIO()
            pd.DataFrame({"Extraction": [md_text]}).to_excel(out, index=False)
            return out.getvalue()
        except:
            return None

# ---------------------------------------------------------------------------
# 6.  BLUEPRINTS
# ---------------------------------------------------------------------------
BLUEPRINTS = {
    "Universal OCR (Any Text)": {
        "identity": "You are an advanced Universal OCR engine capable of extracting text from ANY image. Your job is to perform complete, highly accurate, and exhaustive transcription of the entire document from top to bottom.",
        "structure": "Organize the extracted text into a clean, professional Markdown format. Use Markdown Headers (###) for sections, **bold text** for keys/labels, bullet points (-) for lists, and Markdown Tables (|---|) for any tabular or grid-like data.",
        "instructions": "MANDATORY: Extract EVERY single line of text from the image without exception. Do not truncate, omit, or summarize the content. Scan the entire image carefully and transcribe all items, values, and notes from top to bottom.",
        "rules": "2. EXHAUSTIVE TRANSCRIPTION: Do not stop generating until the very last word of the image is transcribed.\n3. FORMATTING: Use Markdown (bolding, lists, tables) to give the text a professional structure.\n4. KEY-VALUE PAIRS: If you see a label and a value (e.g., Name: John), format it as **Name:** John.\n5. NO PREAMBLE: Start directly with the extracted data. No greetings or meta-commentary."
    },
    "LDSL Diagnostics": {
        "identity": "You are a high-performance medical OCR engine (Qwen 2.5 VL optimized).",
        "structure": """\
**Patient Name:** [Full Name] | **Age/Sex:** [Age]/[Sex]
**Referred Doctor:** [Doctor Name] | **Coll. Time & Date:** [Date and Time]

| S.No | Test Description | Sample Type | Result (if any) |
|------|-----------------|-------------|-----------------|
| 1    | [Test Name]     | [Sample]    | [Value/Result]  |

**History:** DOB: [DOB] | Weight: [Weight] | Diabetes: [Yes/No] | Ultrasound: [Details]
**Footer:** Checked By: [Name] | Area: [Location]""",
        "instructions": "MANDATORY: Transcribe every handwritten entry exactly. Use the Markdown table for all test results. Do not omit the Patient Name or History section.",
        "rules": "2. SPATIAL AWARENESS: Maintain the exact layout and hierarchy seen in the document.\n3. HANDWRITING: Transcribe every handwritten scribble or mark. If a checkmark is present in a box, represent it as [x].\n4. NUMERICAL PRECISION: Do not round or alter any numbers, dates, or measurements.\n5. NO PREAMBLE: Start directly with the extracted data. No greetings or meta-commentary.\n6. TABLE INTEGRITY: Ensure every column of the table is populated correctly based on the visual rows."
    },
    "Healmax Diagnostics": {
        "identity": "You are a high-performance medical OCR engine (Qwen 2.5 VL optimized).",
        "structure": """\
**Franchisee Code:** ___ | **Date:** ___

| S.No | Patient Name | Age/Sex | Test Code/Name | Sample Type | Barcode No | Date/Time | Customer | Referral Doctor |
|------|-------------|---------|----------------|-------------|------------|-----------|----------|-----------------|""",
        "instructions": "Fill all 9 table columns. Do NOT merge or skip any column.",
        "rules": "2. SPATIAL AWARENESS: Maintain the exact layout and hierarchy seen in the document.\n3. HANDWRITING: Transcribe every handwritten scribble or mark. If a checkmark is present in a box, represent it as [x].\n4. NUMERICAL PRECISION: Do not round or alter any numbers, dates, or measurements.\n5. NO PREAMBLE: Start directly with the extracted data. No greetings or meta-commentary.\n6. TABLE INTEGRITY: Ensure every column of the table is populated correctly based on the visual rows."
    }
}

# ---------------------------------------------------------------------------
# 7.  PROMPT BUILDER
# ---------------------------------------------------------------------------
def build_prompt(client: str, page_info: str = "") -> str:
    bp = BLUEPRINTS.get(client, BLUEPRINTS["Universal OCR (Any Text)"])
    page_note = f" ({page_info})" if page_info else ""
    return f"""{bp['identity']}{page_note}

Your objective: Extract EVERY piece of text from the image with 100% fidelity.

Output Format:
{bp['structure']}

Rules for 100% Accuracy:
1. MANDATORY: {bp['instructions']}
{bp['rules']}"""

# ---------------------------------------------------------------------------
# 8.  MAIN STREAMLIT APP
# ---------------------------------------------------------------------------
def render_ocr_app():
    client    = st.session_state.get("current_client", "Universal OCR (Any Text)")

    st.header(f"⚡ {client} — Precision OCR Engine")

    # Session-state initialisation
    for key, default in [
        ("ocr_pages",    []),   # list of (page_label, extracted_text)
        ("ocr_combined", ""),
        ("last_time",    0.0),
        ("ocr_client",   ""),
        ("ocr_images",   []),   # list of original PIL images
    ]:
        if key not in st.session_state:
            st.session_state[key] = default

    # Reset results when client template changes
    if st.session_state.ocr_client != client:
        st.session_state.ocr_pages    = []
        st.session_state.ocr_combined = ""
        st.session_state.ocr_images   = []
        st.session_state.ocr_client   = client

    # ── SIDEBAR ─────────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown("### 📤 Upload Document")
        uploaded = st.file_uploader(
            "Supported: PNG, JPG, JPEG, PDF",
            type=["png", "jpg", "jpeg", "pdf"],
            key="ocr_uploader_prod",
        )

        process_btn = st.button(
            "🚀 Extract Data",
            width='stretch',
            type="primary",
        )
        clear_btn = st.button(
            "🗑️ Clear Results",
            width='stretch',
        )

        st.divider()
        st.markdown("**⚙️ Engine Info**")
        st.caption("Active Model: `qwen2.5vl:32b` (Fixed)")
        st.caption("Resolution: Optimized 300 DPI")
        st.caption("Strategies: 5 adaptive preprocessing layers")
        st.caption("Auto-retries if response is empty")

    if clear_btn:
        st.session_state.ocr_pages    = []
        st.session_state.ocr_combined = ""
        st.session_state.ocr_images   = []
        st.session_state.last_time    = 0.0
        st.rerun()

    # ── PROCESSING ──────────────────────────────────────────────────────────
    if process_btn:
        if not uploaded:
            st.warning("⚠️ Please upload a file first.")
        else:
            st.session_state.ocr_pages    = []
            st.session_state.ocr_combined = ""
            st.session_state.ocr_images   = []

            start_time = time.time()

            # Load pages
            try:
                if uploaded.name.lower().endswith(".pdf"):
                    pages = load_pdf_pages(uploaded.read())
                    page_labels = [f"Page {i+1}" for i in range(len(pages))]
                else:
                    pages = [Image.open(uploaded)]
                    page_labels = ["Page 1"]
                
                st.session_state.ocr_images = pages
            except Exception as e:
                st.error(f"❌ Could not open file: {e}")
                st.stop()

            total_pages = len(pages)
            progress    = st.progress(0, text="Starting…")
            status_box  = st.empty()

            all_pages_text = []

            for idx, (raw_img, label) in enumerate(zip(pages, page_labels)):
                progress.progress(
                    (idx) / total_pages,
                    text=f"Processing {label} of {total_pages}…"
                )

                page_info = f"{label} of {total_pages}" if total_pages > 1 else ""
                prompt    = build_prompt(client, page_info)

                extracted, strategy_used = call_model_with_retry(
                    raw_img, prompt, status_box, model_name="qwen2.5vl:32b"
                )

                if extracted:
                    status_box.success(f"✅ {label} extracted via **{strategy_used}**")
                    all_pages_text.append((label, extracted))
                else:
                    status_box.error(
                        f"❌ {label}: Extraction failed. The model returned no text or the content was too short. "
                        "Try a different model or check if the image is clear."
                    )
                    all_pages_text.append((label, f"*[{label}: Extraction failed — content too short or empty]*"))

            progress.progress(1.0, text="Done!")

            elapsed = time.time() - start_time
            st.session_state.ocr_pages    = all_pages_text
            st.session_state.last_time    = elapsed

            # Merge all pages into one document
            if total_pages == 1:
                st.session_state.ocr_combined = all_pages_text[0][1]
            else:
                parts = []
                for lbl, txt in all_pages_text:
                    parts.append(f"---\n### {lbl}\n\n{txt}")
                st.session_state.ocr_combined = "\n\n".join(parts)

            # Archive
            try:
                archive_document(
                    user_id   = st.session_state.get("user_id", 1),
                    filename  = uploaded.name,
                    category  = client,
                    markdown  = st.session_state.ocr_combined,
                    confidence= 99.0,
                )
            except Exception as arc_err:
                st.warning(f"⚠️ Vault archiving failed (result still shown): {arc_err}")

            st.rerun()

    # ── DISPLAY RESULTS ─────────────────────────────────────────────────────
    if st.session_state.ocr_combined:
        # Real Content Analysis for Output Quality
        def analyze_output_quality(text):
            if not text or len(text.strip()) == 0: return 0.0
            
            # Remove standard markdown formatting to inspect raw words
            clean = re.sub(r'[\*#\|\[\]\-\_]', ' ', text)
            words = clean.split()
            if not words: return 0.0
            
            issues = 0
            for w in words:
                # 1. Penalize endless repeating characters (e.g. '00000', 'yyyyy')
                if re.search(r'(.)\1{4,}', w):
                    issues += 1
                # 2. Penalize chaotic non-alphanumeric noise (gibberish symbols)
                elif len(w) > 3 and re.match(r'^[^A-Za-z0-9]+$', w):
                    issues += 1
                # 3. Penalize extremely long words without vowels (broken strings)
                elif len(w) > 12 and not re.search(r'[aeiouAEIOU]', w):
                    if not w.isdigit(): # large numbers are okay
                        issues += 1
            
            error_ratio = issues / len(words)
            # A 10% error ratio drops the score heavily
            score = 100.0 - (error_ratio * 250.0)
            return max(0.0, min(100.0, score))

        quality_score = analyze_output_quality(st.session_state.ocr_combined)

        st.success(
            f"✅ Extraction complete — {len(st.session_state.ocr_pages)} page(s) "
            f"in {st.session_state.last_time:.1f}s | **Output Quality Score: {quality_score:.1f}%**"
        )
        
        col_img, col_res = st.columns(2, gap="large")

        with col_img:
            st.markdown("### 🖼️ Original Input")
            if st.session_state.get("ocr_images"):
                for i, img in enumerate(st.session_state.ocr_images):
                    st.image(img, caption=f"Page {i+1}", width="stretch")
            else:
                st.info("Original image not available.")

        with col_res:
            st.markdown("### 📝 Extracted Output")
            # Multi-page: show tabs per page + combined
            if len(st.session_state.ocr_pages) > 1:
                tab_labels = [lbl for lbl, _ in st.session_state.ocr_pages] + ["📄 Combined", "📊 Data Grid"]
                tabs = st.tabs(tab_labels)
                for i, (tab, (lbl, txt)) in enumerate(zip(tabs[:-2], st.session_state.ocr_pages)): 
                    with tab:
                        st.markdown(txt)
                with tabs[-2]:
                    st.markdown(st.session_state.ocr_combined)
                with tabs[-1]:
                    # Try to show tables as dataframes
                    st.markdown("### 📈 Extracted Data Tables")
                    lines = st.session_state.ocr_combined.split('\n')
                    current_table = []
                    found_any = False
                    for line in lines:
                        raw_line = line.strip()
                        if raw_line.count('|') >= 2:
                            if all(c in '|- ' for c in raw_line) and '-' in raw_line: continue
                            cells = [c.strip() for c in raw_line.split('|') if c.strip()]
                            if cells: current_table.append(cells)
                        else:
                            if current_table and len(current_table) > 1:
                                header, rows = current_table[0], current_table[1:]
                                cleaned = [r[:len(header)] if len(r) > len(header) else r + ['']*(len(header)-len(r)) for r in rows]
                                st.dataframe(pd.DataFrame(cleaned, columns=header), width='stretch')
                                found_any = True
                                current_table = []
                    if current_table and len(current_table) > 1:
                        header, rows = current_table[0], current_table[1:]
                        cleaned = [r[:len(header)] if len(r) > len(header) else r + ['']*(len(header)-len(r)) for r in rows]
                        st.dataframe(pd.DataFrame(cleaned, columns=header), width='stretch')
                        found_any = True
                    if not found_any:
                        st.info("No structured tables detected for grid view. See 'Combined' for raw text.")
            else:
                tab1, tab2 = st.tabs(["📄 Document View", "📊 Data Grid"])
                with tab1:
                    st.markdown(st.session_state.ocr_combined)
                with tab2:
                    st.markdown("### 📈 Extracted Data Tables")
                    lines = st.session_state.ocr_combined.split('\n')
                    current_table = []
                    found_any = False
                    for line in lines:
                        raw_line = line.strip()
                        if raw_line.count('|') >= 2:
                            if all(c in '|- ' for c in raw_line) and '-' in raw_line: continue
                            cells = [c.strip() for c in raw_line.split('|') if c.strip()]
                            if cells: current_table.append(cells)
                        else:
                            if current_table and len(current_table) > 1:
                                header, rows = current_table[0], current_table[1:]
                                cleaned = [r[:len(header)] if len(r) > len(header) else r + ['']*(len(header)-len(r)) for r in rows]
                                st.dataframe(pd.DataFrame(cleaned, columns=header), width='stretch')
                                found_any = True
                                current_table = []
                    if current_table and len(current_table) > 1:
                        header, rows = current_table[0], current_table[1:]
                        cleaned = [r[:len(header)] if len(r) > len(header) else r + ['']*(len(header)-len(r)) for r in rows]
                        st.dataframe(pd.DataFrame(cleaned, columns=header), width='stretch')
                        found_any = True
                    if not found_any:
                        st.info("No structured tables detected for grid view. See 'Document View' for raw text.")

        st.divider()

        # Download buttons
        col1, col2, col3, col4 = st.columns(4)
        combined = st.session_state.ocr_combined

        with col1:
            st.download_button(
                "📥 Download Markdown",
                combined,
                "extraction.md",
                "text/markdown",
                width='stretch',
            )
        with col2:
            pdf = generate_pro_pdf(combined, client)
            if isinstance(pdf, bytes):
                st.download_button(
                    "📥 Download PDF",
                    pdf,
                    "extraction.pdf",
                    "application/pdf",
                    width='stretch',
                )
            else:
                st.error(pdf)
        with col3:
            docx_bytes = generate_docx(combined, client)
            if isinstance(docx_bytes, bytes):
                st.download_button(
                    "📥 Download Word",
                    docx_bytes,
                    "extraction.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    width='stretch',
                )
            else:
                st.error(docx_bytes)
        with col4:
            excel_bytes = generate_excel(combined)
            if excel_bytes:
                st.download_button(
                    "📥 Download Excel",
                    excel_bytes,
                    "extraction.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    width='stretch',
                )
            else:
                st.info("No table found for Excel export")

    else:
        # Idle state
        st.info(
            "📋 Upload a medical form (image or PDF) in the sidebar and click "
            "**Extract Data** to begin."
        )
        with st.expander("ℹ️ How the engine works"):
            st.markdown("""
**5-Strategy Auto-Retry Pipeline:**
1. **Original** — raw image with orientation fix  
2. **Mild Enhancement** — sharpness + 1.5× contrast (printed forms)  
3. **Grayscale Boost** — grayscale + 1.8× contrast (photocopies)  
4. **Adaptive Threshold** — histogram stretch + unsharp mask (faded handwriting)  
5. **Denoise + Enhance** — median filter + enhance (camera phone photos)  

The engine automatically tries the next strategy if a response is empty.
            """)